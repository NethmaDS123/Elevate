# features/resume_extraction.py

"""
ResumeExtractor: Extracts text content and formatting details from resume files (PDF, DOC, DOCX)
Uses PyPDF2 for PDF files and python-docx for Word documents
"""

import os
import logging
import tempfile
from typing import Dict, Any, Optional
import PyPDF2
from docx import Document
from docx.shared import RGBColor
import io
from fastapi import UploadFile
import re

# For PDF visual analysis
try:
    from PIL import Image
    import fitz  # PyMuPDF
    HAS_VISUAL_LIBS = True
except ImportError:
    HAS_VISUAL_LIBS = False

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s"
)
logger = logging.getLogger("features.resume_extraction")

class ResumeExtractor:
    """Class to handle extraction of text and formatting from resume files"""
    
    def __init__(self):
        pass
        
    async def extract_text(self, file: UploadFile) -> Dict[str, Any]:
        """
        Extract text and formatting details from uploaded resume file
        
        Args:
            file: The uploaded file object
            
        Returns:
            Dict containing extracted text and metadata
        """
        # Get file extension from content type or filename
        file_ext = self._get_file_extension(file)
        
        # Create a temporary file to save the uploaded content
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            # Write uploaded file content to temp file
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        try:
            # Extract text based on file type
            if file_ext.lower() in ['.pdf']:
                text, formatting = self._extract_from_pdf(temp_file_path)
            elif file_ext.lower() in ['.doc', '.docx']:
                text, formatting = self._extract_from_docx(temp_file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
                
            # Get basic metadata
            metadata = {
                "filename": file.filename,
                "content_type": file.content_type,
                "size": len(content),
                "char_count": len(text),
                "formatting": formatting
            }
            
            logger.info(f"Successfully extracted {len(text)} characters from {file.filename}")
            
            return {
                "text": text,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from {file.filename}: {str(e)}")
            raise
        finally:
            # Clean up the temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
    
    def _get_file_extension(self, file: UploadFile) -> str:
        """Determine file extension from content type or filename"""
        content_type_map = {
            "application/pdf": ".pdf",
            "application/msword": ".doc",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx"
        }
        
        # Try to get extension from content type first
        if file.content_type in content_type_map:
            return content_type_map[file.content_type]
        
        # Fall back to filename extension
        if file.filename:
            _, ext = os.path.splitext(file.filename)
            if ext:
                return ext
        
        # Default to .txt if can't determine
        return ".txt"
    
    def _extract_from_pdf(self, file_path: str) -> tuple[str, dict]:
        """Extract text and formatting details from PDF file"""
        text = ""
        formatting = {
            "page_count": 0,
            "has_colors": False,
            "font_sizes": [],
            "line_spacing": "unknown",
            "margins": {},
            "has_images": False,
            "layout_analysis": {},
            "estimated_word_count": 0
        }
        
        # Basic text extraction
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Get page count
            formatting["page_count"] = len(pdf_reader.pages)
            
            # Extract text from each page
            all_text = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                all_text.append(page_text)
                text += page_text + "\n\n"
            
            # Estimate word count
            formatting["estimated_word_count"] = len(re.findall(r'\w+', text))
        
        # Enhanced visual analysis if PyMuPDF is available
        if HAS_VISUAL_LIBS:
            try:
                doc = fitz.open(file_path)
                
                # Check for colors and images
                has_colors = False
                has_images = False
                font_sizes = set()
                
                for page_idx in range(len(doc)):
                    page = doc[page_idx]
                    
                    # Check for images
                    if page.get_images():
                        has_images = True
                    
                    # Extract font information
                    blocks = page.get_text("dict")["blocks"]
                    for block in blocks:
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    if span.get("color") and span["color"] != 0:
                                        has_colors = True
                                    if span.get("size"):
                                        font_sizes.add(round(span["size"]))
                
                formatting["has_colors"] = has_colors
                formatting["has_images"] = has_images
                formatting["font_sizes"] = sorted(list(font_sizes)) if font_sizes else []
                
                # Analyze layout
                if len(doc) > 0:
                    first_page = doc[0]
                    width, height = first_page.rect.width, first_page.rect.height
                    formatting["layout_analysis"] = {
                        "page_width": width,
                        "page_height": height,
                        "aspect_ratio": width / height if height else 0
                    }
                
                doc.close()
            except Exception as e:
                logger.warning(f"Enhanced PDF analysis failed: {str(e)}")
        
        return text.strip(), formatting
    
    def _extract_from_docx(self, file_path: str) -> tuple[str, dict]:
        """Extract text and formatting details from DOCX file"""
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs_text = [para.text for para in doc.paragraphs if para.text.strip()]
        text = "\n\n".join(paragraphs_text)
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    tables_text.append(row_text)
        
        if tables_text:
            text += "\n\n" + "\n".join(tables_text)
        
        # Analyze formatting
        formatting = {
            "page_count": self._estimate_page_count(doc),
            "has_colors": False,
            "font_sizes": [],
            "font_families": [],
            "has_tables": len(doc.tables) > 0,
            "has_headers": False,
            "has_bullets": False,
            "estimated_word_count": 0
        }
        
        # Check for colors, font sizes, and bullets
        font_sizes = set()
        font_families = set()
        has_colors = False
        has_bullets = False
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
                
            # Check for bullet points
            if para.text.strip().startswith(('•', '-', '*', '○', '▪', '▫', '◦')):
                has_bullets = True
            
            # Check for formatting in runs
            for run in para.runs:
                if run.font.size:
                    # Convert to points (half-points to points)
                    size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size / 12700
                    font_sizes.add(round(size_pt))
                
                if run.font.name:
                    font_families.add(run.font.name)
                
                # Check for non-black color
                if run.font.color and run.font.color.rgb and run.font.color.rgb != RGBColor(0, 0, 0):
                    has_colors = True
        
        formatting["has_colors"] = has_colors
        formatting["has_bullets"] = has_bullets
        formatting["font_sizes"] = sorted(list(font_sizes)) if font_sizes else []
        formatting["font_families"] = sorted(list(font_families)) if font_families else []
        formatting["estimated_word_count"] = len(re.findall(r'\w+', text))
        
        # Check for headers/sections
        if doc.sections:
            formatting["has_headers"] = any(section.header.is_linked_to_previous == False for section in doc.sections)
        
        return text.strip(), formatting
    
    def _estimate_page_count(self, doc: Document) -> int:
        """Estimate page count for a Word document"""
        # This is a rough estimation as Word documents don't store page count directly
        # Average characters per page is around 3000 for standard formatting
        total_text = "\n".join([para.text for para in doc.paragraphs])
        char_count = len(total_text)
        
        # Account for tables which take up more space
        table_count = len(doc.tables)
        
        # Very rough estimation
        estimated_pages = max(1, (char_count / 3000) + (table_count * 0.5))
        return round(estimated_pages) 